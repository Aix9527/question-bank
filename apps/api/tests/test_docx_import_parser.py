from __future__ import annotations

import base64
from pathlib import Path

from docx import Document


def build_fixture(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("第一段")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    image_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZQmcAAAAASUVORK5CYII="
    )
    image_path = path.parent / "pixel.png"
    image_path.write_bytes(image_bytes)
    doc.add_paragraph("公式图：").add_run().add_picture(str(image_path))
    doc.add_paragraph("最后一段")
    doc.save(path)
    return path


def test_parser_preserves_body_order_tables_and_images(tmp_path):
    from app.services.imports.docx_parser import parse_docx

    source = build_fixture(tmp_path / "fixture.docx")
    ast = parse_docx(source)

    assert [block["kind"] for block in ast["blocks"]] == ["paragraph", "table", "paragraph", "paragraph"]
    assert ast["blocks"][0]["text"] == "第一段"
    assert ast["blocks"][1]["rows"] == [["A", "B"]]
    assert '<img src="data:image/png;base64,' in ast["blocks"][2]["html"]
    assert ast["media_count"] == 1
    assert len(ast["source_sha256"]) == 64
